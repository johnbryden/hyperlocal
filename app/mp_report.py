"""
MP report — constituency-style .docx from run_process output.

Groups by tag (sum of comments), takes top N tags and up to M posts per tag,
generates MP action suggestions via the same prompt as junk/analyse_mp.py,
and writes a Word report. When a tag has more than M posts, the doc states
that and suggests viewing the output Google Sheet for the full summary.
"""
from __future__ import annotations

import json as json_module
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Union

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from cloudpathlib import AnyPath

import app.ai_wrapper as ai
from app.simple_logger import get_logger

logger = get_logger(__name__)

PathType = Path | AnyPath


def _format_week_starting(date_from: str) -> str:
    """Format date_from (YYYY-MM-DD) as 'Week starting Monday 2nd February 2026'."""
    d = datetime.strptime(date_from, "%Y-%m-%d")
    day = d.day
    if 10 <= day % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    day_ord = f"{day}{suffix}"
    weekday = d.strftime("%A")
    month_year = d.strftime("%B %Y")
    return f"Week starting {weekday} {day_ord} {month_year}"


def _get_prompt_suggest_mp_action(row: pd.Series, constituency_name: str) -> str:
    """Generate prompt for MP action suggestions (same as analyse_mp)."""
    prompt = f"""
    You are a helpful assistant that suggests actions for a local MP based in the {constituency_name} constituency to take based on a specific post.

    The post is: "{row.body}"
    The comments are: "{row.comment_texts}"

    Generate a single sentence summary of the post and commments focussing on the issue being discussed.

    Make a very very brief one-sentence suggestion for very low effort actions (e.g., max 1 hour) that an MP could do about the issue. 
    Focus on visible responses, but also something substantive that actually addresses the issue rather than lip service.
    Act as though you are an MP planning campaign or policy actions. Use the following framework to identify and plan your mechanisms for action. Make sure your suggestions cover all three areas — Parliament, Constituency, and Other — and add any that might be missing.

    Parliament actions (examples):
    Pressure a minister
    Ask a Parliamentary Question (PQ)
    Give a speech in the House
    Hold or join a debate
    Convene an event
    Introduce a Bill
    Ask for a library briefing
    Work through an APPG (All-Party Parliamentary Group)
    Persuade colleagues to support your position

    Constituency actions (examples):
    Contact the Council or other decision maker
    Convene a roundtable
    Write to the local paper
    Create a petition
    Organize a community event (e.g., litter pick, cleanup)
    Post on social media
    Engage local media
    Talk to local groups
    Contact the local police, NHS trust, schools, etc.

    Other (examples):
    Pair up with a national campaign
    Engage with national media

    Respond with a json object with the following keys:
    {json_module.dumps({'summary':'your_summary','mp_action':'your_mp_action'})}
    """
    return prompt


def _is_valid_suggest_mp_action(row: pd.Series) -> bool:
    """Validate MP action suggestion response."""
    if "summary" not in row.index:
        return False
    if "mp_action" not in row.index:
        return False
    return True


def build_tag_report_df(
    df: pd.DataFrame,
    *,
    number_of_tags: int,
    max_posts_per_tag: int,
) -> pd.DataFrame:
    """
    Group by tag, sum comments, take top number_of_tags, then up to
    max_posts_per_tag posts per tag (by comment count). Returns a single
    dataframe with columns total_comments, num_posts merged onto each row.
    """
    if "tag" not in df.columns or "comments" not in df.columns:
        return pd.DataFrame()

    grouped = (
        df.groupby("tag")
        .agg(
            total_comments=pd.NamedAgg(column="comments", aggfunc="sum"),
            num_posts=pd.NamedAgg(column="comments", aggfunc="count"),
        )
        .sort_values("total_comments", ascending=False)
    )
    # Drop sentinel if present
    if "Non-specific" in grouped.index:
        grouped = grouped.drop(index="Non-specific", errors="ignore")
    top_tags = grouped.head(number_of_tags).reset_index()

    # For each tag, take top max_posts_per_tag posts by comments
    rows = []
    for _, tag_row in top_tags.iterrows():
        tag_name = tag_row["tag"]
        total_comments = tag_row["total_comments"]
        num_posts = tag_row["num_posts"]
        tag_posts = (
            df[df["tag"] == tag_name]
            .nlargest(max_posts_per_tag, "comments")
            .copy()
        )
        tag_posts["total_comments"] = total_comments
        tag_posts["num_posts"] = num_posts
        rows.append(tag_posts)
    if not rows:
        return pd.DataFrame()
    return pd.concat(rows, ignore_index=True)


def generate_mp_suggestions_for_df(
    df: pd.DataFrame,
    constituency_name: str,
    model: str,
) -> pd.DataFrame:
    """Run MP suggestion AI on each row; returns dataframe with summary, mp_action."""
    def get_prompt(row: pd.Series) -> str:
        return _get_prompt_suggest_mp_action(row, constituency_name)

    return ai.iterate_df_rows(
        df,
        get_prompt=get_prompt,
        is_valid=_is_valid_suggest_mp_action,
        response="dict",
        model=model,
    )


def _add_hyperlink(paragraph, url: str, text: str, *, font_name: str | None = None):
    """Add a hyperlink to a paragraph. Optionally set font_name for the link run."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font_name:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_constituency_report(
    report_df: pd.DataFrame,
    output_path: PathType,
    constituency_name: str,
    date_from: str,
    date_to: str,
    *,
    max_posts_per_tag: int,
    total_posts: int | None = None,
    n_political: int = 0,
    output_sheet_url: str | None = None,
) -> tuple[PathType, bytes]:
    """
    Write a Word report: one section per tag with stats and a table of posts
    (summary, link, suggested MP action). If a tag has more than max_posts_per_tag
    posts, add a note to view the output Google Sheet for the full summary.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Constituency Report: {constituency_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_para.add_run(_format_week_starting(date_from))
    run.font.size = Pt(12)
    run.font.italic = True

    if total_posts is not None and total_posts > 0:
        pct = round(100 * n_political / total_posts) if total_posts else 0
        summary_text = (
            f"We found {total_posts} posts for the week, of which {n_political} ({pct}%) "
            "were classified as political."
        )
    else:
        summary_text = (
            f"We found {n_political} posts classified as political for the week."
        )
    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_para.add_run(summary_text).font.size = Pt(11)

    doc.add_paragraph()

    doc.add_heading("Tags Summary", level=1)

    for tag_name in report_df["tag"].unique():
        tag_block = report_df[report_df["tag"] == tag_name]
        if tag_block.empty:
            continue
        first = tag_block.iloc[0]
        total_comments = int(first["total_comments"])
        num_posts = int(first["num_posts"])

        doc.add_heading(tag_name, level=2)
        stats_para = doc.add_paragraph()
        stats_para.add_run(f"Comments: {total_comments}  •  Posts: {num_posts}")

        if num_posts > max_posts_per_tag:
            note = doc.add_paragraph()
            note.add_run(
                f"More than {max_posts_per_tag} posts in this tag. "
                "View the output Google Sheet for the full summary."
            ).font.italic = True

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Post Summary"
        header_cells[1].text = "Link"
        header_cells[2].text = "Suggested MP Action"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E7E6E6")
            cell._tc.get_or_add_tcPr().append(shading)

        for _, post in tag_block.iterrows():
            row_cells = table.add_row().cells
            summary = post.get("summary", "No summary available")
            row_cells[0].text = str(summary) if pd.notna(summary) else "No summary available"
            url = post.get("url", "")
            if pd.notna(url) and url:
                link_para = row_cells[1].paragraphs[0]
                link_para.add_run("  ")
                _add_hyperlink(link_para, str(url), "View Post")
            else:
                row_cells[1].text = "No link"
            mp_action = post.get("mp_action", "No action suggested")
            row_cells[2].text = str(mp_action) if pd.notna(mp_action) else "No action suggested"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        for row in table.rows:
            row.cells[0].width = Inches(3.0)
            row.cells[1].width = Inches(1.0)
            row.cells[2].width = Inches(2.5)
        doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.add_run("Full data for this week is available in the output spreadsheet: ")
    if output_sheet_url:
        _add_hyperlink(footer_para, output_sheet_url, "View data sheet")
    else:
        footer_para.add_run("(link not available)").font.italic = True

    producer_para = doc.add_paragraph()
    producer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    producer_para.add_run("Produced by ")
    _add_hyperlink(producer_para, "https://categorum.ai", "categorum.ai", font_name="Consolas")
    producer_para.add_run(" in association with ")
    _add_hyperlink(producer_para, "https://campaignlab.uk/", "Campaign Lab")

    # python-docx only supports local paths; write via buffer for GCS compatibility
    buf = BytesIO()
    doc.save(buf)
    doc_bytes = buf.getvalue()
    path = AnyPath(output_path)
    with path.open("wb") as f:
        f.write(doc_bytes)

    logger.info(
        "MP report saved.",
        extra={"output_path": str(output_path), "constituency": constituency_name},
    )
    return output_path, doc_bytes


def generate_report(
    df: pd.DataFrame,
    output_path: PathType,
    constituency_name: str,
    date_from: str,
    date_to: str,
    *,
    number_of_tags: int,
    max_posts_per_tag: int,
    model: str,
    total_posts: int | None = None,
    n_political: int | None = None,
    output_sheet_url: str | None = None,
) -> tuple[PathType, bytes] | None:
    """
    Build tag report dataframe, generate MP suggestions, and write .docx.
    Returns (output_path, docx_bytes) or None if no data. Caller may use
    docx_bytes to upload to Google Drive.
    """
    report_df = build_tag_report_df(
        df,
        number_of_tags=number_of_tags,
        max_posts_per_tag=max_posts_per_tag,
    )
    if report_df.empty:
        logger.info(
            "No tags or posts for MP report; skipping.",
            extra={"constituency": constituency_name, "date_from": date_from, "date_to": date_to},
        )
        return None

    if n_political is None:
        n_political = len(df)

    with_suggestions = generate_mp_suggestions_for_df(
        report_df,
        constituency_name=constituency_name,
        model=model,
    )
    _, doc_bytes = create_constituency_report(
        with_suggestions,
        output_path,
        constituency_name,
        date_from,
        date_to,
        max_posts_per_tag=max_posts_per_tag,
        total_posts=total_posts,
        n_political=n_political,
        output_sheet_url=output_sheet_url,
    )
    return output_path, doc_bytes
